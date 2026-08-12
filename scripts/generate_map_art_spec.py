#!/usr/bin/env python3
"""Generate and validate the first formal map art specification."""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables" / "03_统一地图美术规范.docx"
VERSION = "V0.1"
DOC_ID = "MAP-ART-SPEC-001"
EXCEL_NAME = "02_地图坐标与对象锚点表.xlsx"

INK = "263238"
ACCENT = "8A5A2B"
PALE = "F3EBDD"
MIST = "EEF1EE"
LINE = "B7AA92"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total = int(round(sum(widths_cm) / 2.54 * 1440))
    widths = [int(round(w / 2.54 * 1440)) for w in widths_cm]
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, east_asia="Microsoft YaHei", latin="Arial", size=None, bold=None, color=None, italic=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key, value in (("ascii", latin), ("hAnsi", latin), ("eastAsia", east_asia)):
        rfonts.set(qn(f"w:{key}"), value)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size, bold=False, color=INK, east_asia="Microsoft YaHei", latin="Arial"):
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)


def add_field(paragraph, instruction: str, display="1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_page_number(paragraph) -> None:
    add_field(paragraph, " PAGE ", "1")


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.first_line_indent = Cm(0.74)

    specs = {
        "Title": (28, True, INK, 0, 12),
        "Subtitle": (13, False, ACCENT, 0, 8),
        "Heading 1": (17, True, INK, 14, 7),
        "Heading 2": (13.5, True, ACCENT, 10, 5),
        "Heading 3": (11.5, True, INK, 8, 3),
        "Caption": (9, False, "666666", 4, 8),
        "Quote": (10.5, False, ACCENT, 6, 8),
    }
    for name, (size, bold, color, before, after) in specs.items():
        style = styles[name]
        set_style_font(style, size, bold, color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")
        style.paragraph_format.first_line_indent = Cm(0)

    aliases = {
        "标题": "Title",
        "副标题": "Subtitle",
        "标题1": "Heading 1",
        "标题2": "Heading 2",
        "标题3": "Heading 3",
        "正文": "Normal",
        "引用": "Quote",
        "图题": "Caption",
        "表题": "Caption",
    }
    for custom, base in aliases.items():
        if custom not in styles:
            style = styles.add_style(custom, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles[base]

    for custom, fill, color in (("警告", "FFF2CC", "7F6000"), ("规则", MIST, INK)):
        if custom not in styles:
            style = styles.add_style(custom, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles["Normal"]
            set_style_font(style, 10, True, color)
            style.paragraph_format.left_indent = Cm(0.55)
            style.paragraph_format.right_indent = Cm(0.35)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(7)
        styles[custom]._custom_fill = fill

    for style_name in ("List Bullet", "List Number"):
        set_style_font(styles[style_name], 10.3)
        styles[style_name].paragraph_format.space_after = Pt(2)
        styles[style_name].paragraph_format.line_spacing = 1.15


def apply_paragraph_fill(paragraph, fill: str, border=LINE) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), border)
    p_bdr.append(bottom)
    ppr.append(p_bdr)


def add_rule(doc, text: str, warning=False):
    p = doc.add_paragraph(style="警告" if warning else "规则")
    apply_paragraph_fill(p, "FFF2CC" if warning else MIST)
    p.add_run(text)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if level:
            p.paragraph_format.left_indent = Cm(1.2 + level * 0.5)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows, widths, caption=None):
    if caption:
        p = doc.add_paragraph(caption, style="表题")
        p.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(text)
        set_cell_shading(cell, PALE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9, bold=True, color=INK)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_repeat_table_header(table.rows[0])
    for r, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            if r % 2:
                set_cell_shading(cells[idx], "FAF9F6")
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.first_line_indent = Cm(0)
                for run in p.runs:
                    set_run_font(run, size=8.7)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_figure_placeholder(doc, number, title):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [16.1])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8F6F0")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("待项目确认图")
    set_run_font(run, size=12, bold=True, color="777777")
    cap = doc.add_paragraph(f"图{number} {title}", style="图题")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_running_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    hp = section.header.paragraphs[0]
    hp.text = f"统一地图美术规范 {VERSION}"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        set_run_font(run, size=8.5, color="6E6E6E")
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(f"{DOC_ID}  |  第 ")
    set_run_font(r, size=8.5, color="6E6E6E")
    add_page_number(fp)
    r = fp.add_run(" 页")
    set_run_font(r, size=8.5, color="6E6E6E")


def add_heading(doc, level, text):
    return doc.add_paragraph(text, style=f"标题{level}")


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    set_update_fields(doc)
    core = doc.core_properties
    core.title = "统一地图美术规范"
    core.subject = "Unified World Map Art Specification"
    core.author = "MandateOfHeroes 项目组"
    core.keywords = "地图, 美术规范, LOD, 坐标, 锚点, 中山"
    core.comments = f"{DOC_ID}; {VERSION}; generated by scripts/generate_map_art_spec.py"

    cover = doc.sections[0]
    cover.page_width, cover.page_height = Cm(21), Cm(29.7)
    cover.top_margin, cover.bottom_margin = Cm(2.4), Cm(2.2)
    cover.left_margin = cover.right_margin = Cm(2.45)
    cover.different_first_page_header_footer = True
    cover.header_distance, cover.footer_distance = Cm(1.2), Cm(1.2)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("统一地图美术规范")
    set_run_font(r, size=28, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Unified World Map Art Specification")
    set_run_font(r, size=14, color=ACCENT)
    doc.add_paragraph()
    meta = [
        ("版本", VERSION), ("适用阶段", "东汉天下母版与中山连续地图样板"),
        ("文件编号", DOC_ID), ("对应数据模板", EXCEL_NAME),
        ("状态", "第一版／待项目确认"), ("生成日期", date.today().isoformat()),
    ]
    add_table(doc, ["项目", "内容"], meta, [4.0, 12.1])
    p = doc.add_paragraph("同一世界，不同比例尺；地理不重置，细节逐级展开。", style="引用")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    add_rule(doc, "历史与地理警示：未经确认的历史坐标、城市位置与行政关系均为待考证信息，不得写入最终项目事实。", True)

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width, body.page_height = Cm(21), Cm(29.7)
    body.top_margin, body.bottom_margin = Cm(2.2), Cm(2.0)
    body.left_margin = body.right_margin = Cm(2.45)
    body.header_distance, body.footer_distance = Cm(1.0), Cm(1.0)
    set_running_header_footer(body)

    add_heading(doc, 1, "文档控制")
    add_heading(doc, 2, "文档信息")
    add_table(doc, ["项目", "内容"], [
        ("文档名称", "统一地图美术规范"), ("英文名称", "Unified World Map Art Specification"),
        ("文档编号", DOC_ID), ("当前版本", VERSION), ("所属系统", "统一世界地图系统"),
        ("适用范围", "天下战略图至建筑院落图"), ("数据源", EXCEL_NAME), ("文档状态", "待项目确认"),
    ], [4.0, 12.1], "表1 文档信息")
    add_heading(doc, 2, "版本记录")
    add_table(doc, ["版本", "日期", "修改人", "修改章节", "修改说明", "审核人", "状态"],
              [(VERSION, date.today().isoformat(), "", "全部", "第一版生成", "", "待确认"), ("", "", "", "", "", "", "")],
              [1.5, 2.1, 2.0, 2.2, 3.7, 2.0, 2.6], "表2 版本记录")
    add_heading(doc, 2, "审核签署")
    add_table(doc, ["角色", "姓名／签字", "日期", "意见"],
              [(x, "", "", "") for x in ("策划负责人", "历史资料负责人", "地图负责人", "美术负责人", "技术美术", "客户端负责人", "项目负责人")],
              [4.2, 4.0, 3.1, 4.8], "表3 审核签署")

    doc.add_page_break()
    add_heading(doc, 1, "目录")
    toc = doc.add_paragraph()
    add_field(toc, ' TOC \\o "1-3" \\h \\z \\u ', "打开 Word 后更新目录域")
    add_rule(doc, "目录为自动域。首次打开文档后，请右键目录并选择“更新域 → 更新整个目录”。")

    add_heading(doc, 1, "1. 文档目的与适用范围")
    doc.add_paragraph("本规范不是单张插画规范，而是地图概念设计、多尺度空间设计、场景模块设计、游戏对象资产规范、LOD表现规范与跨层级连续性规范的综合标准。它面向策划、史料、地图、美术、UI、技术美术、程序、数据工具和测试岗位。")
    add_bullets(doc, ["天下战略图、州郡区域图与县域乡野图", "城镇空间图、村庄／庄园／坞堡图", "建筑院落与室内图、战场地图、专题数据覆盖图"])
    add_rule(doc, "状态声明：连续比例尺与完整专题视图属于已定方案；现有地图实现仍为原型或最小底座。本规范定义目标合同，不代表全部功能已经实现。")

    add_heading(doc, 1, "2. 项目核心原则")
    add_numbered(doc, [
        "所有层级引用同一世界坐标；所有对象拥有唯一稳定对象ID。",
        "放大只增加细节，不改变基础位置；河流、道路、城市、城门和建筑跨层连续。",
        "不得为各层级建立互不关联的地图；战场继承原县域、村庄或城镇地形。",
        "战斗造成的桥梁、建筑、道路损毁必须反馈到原世界，并持续到修复。",
        "情报显示不等于世界真实数据；专题模式只增加受权限约束的数据覆盖。",
        "城市、旗帜、军队、人物、库存、所有权与损毁不得永久画死在底图。",
        "动态对象必须可由程序独立加载、隐藏、替换；美术不得自行修改已确认坐标。",
    ])

    add_heading(doc, 1, "3. 世界地图层级体系")
    levels = [
        ("L01", "天下战略图", "天下局势、长距离移动", "山河、州域、重要城市、军团、战略道路", "普通县村、农田、建筑"),
        ("L02", "州郡区域图", "郡县与交通网络", "郡治、县城、关隘、港口、官道", "普通房屋和人物"),
        ("L03", "县域乡野图", "地方经营与探索", "村庄、庄园、农田、水利、资源设施", "室内库存"),
        ("L04", "城镇空间图", "城内移动与活动", "城墙、街道、坊区、实体建筑、居民", "天下战略信息"),
        ("L05", "村庄庄园图", "家庭与基层生产", "房屋、田地、牲畜、仓储、水井、作坊", "远距离军团"),
        ("L06", "建筑院落图", "建筑功能使用", "院落、房间、人物、工作点、可见库存", "州郡级信息"),
        ("L07", "战场地图", "战术战斗", "原有地形、部队、工事、火灾、损毁", "无关经营数据"),
    ]
    add_table(doc, ["编码", "层级", "玩家关注", "必须显示", "不应显示"], levels, [1.3, 2.4, 3.0, 5.8, 3.6], "表4 世界地图层级")
    doc.add_paragraph("L01至L06按比例尺逐级展开：缩小时隐藏低层对象，放大时增加真实细节。L07不是陌生副本，而是L03、L04或L05在战术状态下的展开。")
    add_figure_placeholder(doc, 1, "统一地图层级关系示意")

    doc.add_page_break()
    add_heading(doc, 1, "4. 统一世界坐标与对象ID")
    doc.add_paragraph("世界逻辑坐标独立于画布分辨率。PSB画布调整不得修改对象世界坐标；局部地图必须登记世界范围。城市、建筑、道路、河流通过稳定ID关联，路径保存节点或矢量数据，禁止依靠画师目测重新布置。建筑同时登记世界坐标、城镇局部坐标与院落局部坐标。")
    add_rule(doc, f"数据权威：对象ID、世界坐标、父子关系、锚点、LOD登记和状态编码以 {EXCEL_NAME} 为准；美术分层、视觉表现和资产输出以本规范为准。两者必须通过版本号关联。")
    add_bullets(doc, ["01_世界设置：坐标系、边界和画布换算参数", "02_固定地理、03_地点锚点、04_路径网络：世界事实与连接", "05_层级表现LOD、06_建筑锚点、08_资产清单、09_连续性验收：表现与交付关联"])
    p = doc.add_paragraph()
    r = p.add_run("像素X = 世界X ÷ 世界宽度 × 画布宽度\n像素Y = 世界Y ÷ 世界高度 × 画布高度")
    set_run_font(r, east_asia="Microsoft YaHei", latin="Consolas", size=10, color=ACCENT)
    doc.add_paragraph("公式仅用于显示换算，不改变世界坐标。局部裁切还需应用局部世界边界的原点偏移和尺度参数。")
    add_figure_placeholder(doc, 5, "世界坐标与局部坐标换算示意")

    add_heading(doc, 1, "5. 固定地理与动态对象分离")
    add_heading(doc, 2, "5.1 固定地理")
    add_bullets(doc, ["海岸线、海域、主要山系、河流与湖泊", "平原、盆地、丘陵、高原、河谷、峡谷和天然关口", "基础土地与气候分区"])
    add_heading(doc, 2, "5.2 动态对象")
    add_bullets(doc, ["城市、建筑、城墙、城门、道路、桥梁、港口与渡口", "农田、生产设施、人物、军队、商队、旗帜与所有权", "施工、火灾、损毁、围城、封锁、前线、资源发现与情报标记"])
    add_rule(doc, "即使城市具有历史固定位置，也不得与不可编辑底图合并：城市可能扩建、损坏、易主、改名或废弃。装饰全景不得成为设施数量、权限、价格或库存的世界事实。", True)
    add_figure_placeholder(doc, 2, "固定地理与动态对象分层示意")

    doc.add_page_break()
    add_heading(doc, 1, "6. 天下战略图美术规范")
    doc.add_paragraph("层级定位为天下局势、战略行军和长距离运输。美术方向采用原创的“东汉军府舆图＋绢本设色山水＋高可读性战略地图”，视角为高角度斜俯视。")
    add_heading(doc, 2, "6.1 显示与交互")
    add_bullets(doc, ["必须显示海岸线、主要水系、太行山／秦岭／大别山／南岭等大型山系，以及平原、盆地、丘陵、荒地和大型河谷。", "主要城市、关隘、港口、战略道路和大型军团按情报权限显示。", "山脉限制行军、河流影响运输、关口控制道路、平原支持农业与军团活动。", "点击战略对象进入区域层；进入点必须保持同一对象ID、位置与连接关系。"])
    add_heading(doc, 2, "6.2 禁止项与验收")
    add_bullets(doc, ["不得显示每个县、村庄、农田、矿脉、普通建筑、普通人物或精确库存价格。", "不得复制商业游戏的地图画面、笔触、UI、图标、字体或资产。", "验收以山河战略可读性、主要路径连续性、情报不泄露和缩放锚点一致为准。"])

    add_heading(doc, 1, "7. 州郡区域图美术规范")
    add_bullets(doc, ["显示州、郡国和县的空间关系，以及郡治、国治、县城、关隘、港口、渡口、驿站和大型坞堡。", "显示官道、商路、山路、水运、支流、山岭、河谷，以及大型农耕／牧业／林业／矿业潜力区。", "地方驻军、战略设施、战争前线和危险区受情报门控。", "郡治视觉规模大于县城；关隘跨越真实道路或山口；港口与水面相连；道路粗细表达等级而非装饰。"])

    add_heading(doc, 1, "8. 县域乡野图美术规范")
    add_bullets(doc, ["显示县城城门方向、乡亭里村、庄园坞堡、农田类型、水利、资源与生产设施。", "显示官道、村路、小径、桥梁、渡口、废墟和隐秘地点，以及施工和废弃状态。", "人物、商队、军队、差役、流民按对象或聚合规则显示。", "未知区域仅显示基础地貌；未发现资源不出现设施；过期信息褪色；疑似目标只显示范围。"])

    doc.add_page_break()
    add_heading(doc, 1, "9. 城镇空间图美术规范")
    add_rule(doc, "城镇不是建筑功能菜单的背景。可进入建筑必须真实存在于街区空间中，并读取持久设施事实。")
    add_bullets(doc, ["城墙、城门、角楼、护城河、主干道、巷道、桥梁与排水沟渠。", "官署、市场、居住、手工业、军营等街区，以及城外市集、寺观、墓地和农田。", "城外道路、河流沟渠、城门方向与县域图一致；主要建筑街区位置固定。", "点击建筑前可见真实门面和入口；开放、受限、停运和已进入等状态由世界数据决定。"])

    add_heading(doc, 1, "10. 村庄、庄园与坞堡规范")
    add_heading(doc, 2, "10.1 村庄")
    add_bullets(doc, ["家庭住宅、粮仓、柴房、牲畜棚、水井、池塘、磨坊、晒场。", "田地、菜地、桑园、果园、村祠、亭舍、村路、田埂、沟渠。", "居民、牲畜、农具以及荒田、空屋和灾害覆盖状态。"])
    add_heading(doc, 2, "10.2 豪族庄园与坞堡")
    add_bullets(doc, ["围墙、门楼、瞭望设施、主宅、客舍、仓库、账房。", "佃户区、工匠区、私兵训练场、武器库、水利设施、祠堂、族学和墓地。", "房屋、田地和仓库对应数据对象；家庭数量、人口和库存不得画死。"])

    add_heading(doc, 1, "11. 建筑院落与室内规范")
    doc.add_paragraph("通用空间至少包含建筑外观或院门、前后院、主要房间、在场人物、工作点、可见库存、关联设施入口和运营状态。")
    add_heading(doc, 2, "11.1 中山商号主堂")
    add_bullets(doc, ["门面、独立招牌、会客区、掌柜与账房桌。", "账簿、算盘、契约、地图、货样、任务或告示区。", "仓库与车马场入口、商号成员与访客；入口朝向和相对位置继承街区锚点。"])
    add_heading(doc, 2, "11.2 仓库")
    add_bullets(doc, ["分区货架、装卸区、看守、搬运工、防火防潮设施和门禁。", "粮食、布帛、盐铁和药材按类型表现；不足、堆满、损坏与失窃用状态覆盖，不绘制精确数量。"])

    doc.add_page_break()
    add_heading(doc, 1, "12. 战场地图继承规范")
    add_bullets(doc, ["野外战场继承山坡、树林、河流、道路、村庄、农田、桥梁、渡口和高地。", "城市攻防继承城墙、城门、护城河、城楼、城内道路、关键建筑、仓库和居民区。", "部队、军旗、阵型、辎重、营寨、工事、攻城器械、火势、烟尘、尸体与损毁为战斗动态层。"])
    add_rule(doc, "战场中的桥就是原世界中的桥。桥毁、建筑火灾和道路阻断必须回写同一世界账；战斗结束后不得自动消失。")

    add_heading(doc, 1, "13. 专题地图覆盖规范")
    overlays = [
        ("行政", "辖区、驻地、权属、权限", "低饱和边界色，30%—45%透明", "行政认知与授权"),
        ("军事", "军队、补给、威胁、工事", "克制红褐／蓝灰，危险优先", "带时间的观察快照"),
        ("商业", "市场、价格区间、商路", "赭黄与青绿，避免显示精确未知值", "个人或组织商业情报"),
        ("资源生产", "已知资源、设施、产能状态", "图标＋半透明热区", "勘探与生产记录"),
        ("人口民生", "人口压力、流民、灾情", "连续色带，限制同屏标签", "户籍、调查或传闻"),
        ("建设", "项目、选址、施工、阻塞", "轮廓与进度覆盖", "建设委任权限"),
        ("情报", "未知、疑似、过期、较新", "雾、虚线、褪色、正常", "来源、时间、可信度"),
    ]
    add_table(doc, ["视图", "覆盖信息", "色彩／透明度", "权限"], overlays, [2.1, 5.0, 5.0, 4.0], "表5 专题地图覆盖")
    doc.add_paragraph("专题视图必须复用同一底图和实体，只做筛选、聚合、标注与授权显示。标签采用优先级、避让、聚合和逐级展开；不得为专题模式重绘另一张地图。")

    add_heading(doc, 1, "14. 情报、探索与未知状态")
    add_table(doc, ["状态", "含义", "视觉表现"], [
        ("真实存在但未知", "主体尚未获得信息", "薄雾、留白或绢纸纹理"),
        ("疑似存在", "仅有模糊线索", "大致范围、不确定轮廓"),
        ("已知但过期", "曾掌握但可能变化", "褪色、模糊、显示时间"),
        ("已知且较新", "当前可信度较高", "清晰轮廓、正常饱和度"),
    ], [4.0, 5.7, 6.4], "表6 情报状态")
    add_bullets(doc, ["未发现矿脉不显示矿洞；疑似匪患不显示精确寨门。", "UI显示信息来源、观察时间与可信度；情报迷雾为独立图层。", "信息显示不改变世界对象；视图不得绕过个人知识、组织授权与时效。"])

    doc.add_page_break()
    add_heading(doc, 1, "15. 地形、水系与道路标准")
    add_heading(doc, 2, "15.1 山地")
    doc.add_paragraph("天下图显示山系走向，州郡图显示山岭河谷，县域图显示坡地树林与通行口，战场图显示高低差和遮挡。山体走向不得跨层级改变。")
    add_heading(doc, 2, "15.2 河流")
    doc.add_paragraph("河流中心线跨层一致；放大时增加河岸、浅滩、支流、桥梁和渡口。主河道流向不得改变，港口必须临水；宽度可细化但拓扑不变。")
    add_heading(doc, 2, "15.3 道路")
    add_bullets(doc, ["结构化字段：等级、宽度、容量、速度系数、季节与天气影响、控制权、施工和损坏状态。", "美术通过路宽、颜色、车辙、桥梁等级与维护状态表达等级。", "道路参考层与正式游戏道路资产分离，禁止以装饰线替代路径网络。"])

    add_heading(doc, 1, "16. 城市、建筑和设施标准")
    add_table(doc, ["层级", "城市表现"], [("天下图", "抽象城市符号"), ("州郡图", "城池轮廓"), ("县域图", "城墙、城门和城外道路"), ("城镇图", "完整街区和实体建筑"), ("建筑图", "院落和房间")], [4.0, 12.1], "表7 城市LOD")
    add_bullets(doc, ["建议图层：主体、屋顶、入口、招牌、旗帜、货物、人物、阴影、所有权、施工、损坏、火灾、废弃。", "文字、势力旗帜、所有者颜色、商品数量、人物、火灾、损毁和围城不得永久画入主体。"])

    add_heading(doc, 1, "17. 人物、军队、车辆和商队标准")
    add_bullets(doc, ["对象含玩家、军团、斥候、商队、运输队、流民、舰队、差役、居民、牲畜和车辆。", "天下图显示军团或大队；州郡显示军队、商队、舰队；县域显示队伍与主要人物；城镇／村庄显示具体人物车辆牲畜；院落显示角色与工作状态。", "单位使用独立阴影和标准锚点；势力颜色不进入主体；聚合不改变实际单位坐标；位置必须符合道路、河流与地形。"])

    add_heading(doc, 1, "18. 建设、损坏、火灾和废弃状态")
    states = ["正常", "关闭", "建设中", "扩建", "损坏", "火灾", "废弃", "被占用", "被围困", "被淹", "被劫掠", "未知"]
    add_table(doc, ["状态", "主体／覆盖建议", "系统影响登记"], [(s, "优先保留主体；脚手架、火焰、烟雾、焦土、碎石等使用独立覆盖", "碰撞、通行、入口、功能、人口、库存、所有权按07_动态状态登记") for s in states], [2.4, 7.4, 6.3], "表8 动态状态表现矩阵")
    add_figure_placeholder(doc, 6, "建筑主体与状态覆盖层示意")

    doc.add_page_break()
    add_heading(doc, 1, "19. PSB／PSD分层规范")
    layer_tree = [
        ("00_GUIDE", "辅助", "world_boundary, world_grid, zoom_regions, scale_reference"),
        ("01_SKY_DISTANCE", "远景", "sky, distant_cloud, distant_mountain"),
        ("02_TERRAIN_BASE", "地形底色", "base_land, plateau, basin, wasteland"),
        ("03_MOUNTAIN", "山体", "major_mountain, ridge, hill, mountain_shadow"),
        ("04_PLAIN_LAND", "平原土地", "plain, farmland_reference, grassland, wetland"),
        ("05_WATER", "水系", "sea, lake, major_river, tributary, water_shadow"),
        ("06_VEGETATION", "植被", "north_forest, south_forest, grass, marsh"),
        ("07_ROAD_FIXED_REFERENCE", "道路参考", "major_road_reference, secondary_road_reference, water_route_reference, bridge_reference"),
        ("08_FIXED_LANDMARK", "固定地标", "natural_pass, canyon, island, fixed_rock"),
        ("09_DYNAMIC_BUILDING", "动态建筑", "city, pass, port, bridge, building"),
        ("10_PERSON_UNIT", "人物单位", "player, army, caravan, civilian, vehicle, ship"),
        ("11_FLAG_OWNER", "旗帜权属", "flag, owner_color, territory_overlay"),
        ("12_WEATHER_SEASON", "天气季节", "rain, snow, fog, seasonal_tone"),
        ("13_FIRE_DAMAGE", "火灾损毁", "fire, smoke, rubble, scorched_land, flood"),
        ("14_INTELLIGENCE_FOG", "情报迷雾", "unknown, suspected, expired, known"),
        ("15_UI_REFERENCE", "UI参考", "selection, highlight, zoom_frame, label_reference"),
    ]
    add_table(doc, ["组名", "中文用途", "子层"], layer_tree, [4.1, 3.0, 9.0], "表9 PSB／PSD标准图层树")
    add_bullets(doc, ["辅助层不导出；动态组可独立关闭；阴影尽量独立；文字不进入底图。", "关键矢量路径不得随意栅格化；智能对象保留源文件；参考路径与正式道路资产明确区分。"])
    add_figure_placeholder(doc, 4, "PSB标准图层树")

    add_heading(doc, 1, "20. 透明对象与资产输出规范")
    add_table(doc, ["项目", "标准"], [
        ("预览确认稿", "3840×2160 PNG"), ("天下母版", "PSB，建议15360×8640或更高；最终值待确认"),
        ("动态对象", "透明PNG"), ("色彩空间", "sRGB"), ("主体位深", "16位"),
        ("引擎导出位深", "按引擎要求"), ("透明边缘", "不得有白边或黑边"),
        ("阴影", "优先独立输出"), ("文字", "独立UI层"), ("锚点", "登记到Excel"), ("版本", "可追踪"),
    ], [5.0, 11.1], "表10 资产输出标准")
    add_bullets(doc, ["统一PNG裁切与锚点规则；不以不透明背景替代透明对象。", "输出后检查透明边缘、尺寸、锚点与Excel资产清单的一致性。"])

    add_heading(doc, 1, "21. 文件命名和版本管理")
    p = doc.add_paragraph()
    examples = "\n".join(["terrain_north_plain_base_v001.psb", "river_yellow_main_v001.png", "city_zhongshan_normal_l01_v001.png", "city_zhongshan_damaged_l02_v001.png", "building_shop_zhongshan_normal_l04_v001.png", "overlay_fire_medium_v001.png", "overlay_intelligence_expired_v001.png"])
    r = p.add_run(examples)
    set_run_font(r, latin="Consolas", size=9.5, color=ACCENT)
    add_bullets(doc, ["小写英文、下划线、无空格；版本使用三位数字。", "禁止“最终版”“最新版”“新建文件”等不可追踪名称。", "废弃资产不直接删除，须在资产表登记；每项正式资产具有唯一资产ID，并可追溯对象ID、状态与LOD。"])

    doc.add_page_break()
    add_heading(doc, 1, "22. LOD与缩放切换规则")
    add_bullets(doc, ["每个对象登记最低和最高显示层级；城市由图标切换为轮廓和实体城镇。", "道路逐级增加材质与边缘细节；河流增加河岸、浅滩和支流；单位允许聚合。", "标签采用优先级、避让和聚合；切换可淡入淡出。", "LOD不得修改世界坐标、道路连接或对象父子关系。"])

    add_heading(doc, 1, "23. 光照、色彩和视角标准")
    add_table(doc, ["项目", "第一版建议"], [
        ("视觉方向", "古代绢本设色山水与战略地图结合"), ("视角", "高角度斜俯视"),
        ("默认光照", "左上至右下"), ("整体色彩", "土黄、黛青、灰绿、河水青灰"),
        ("天下层级", "低对比，突出山河走势"), ("局部层级", "逐步增加材质、阴影与生活细节"),
    ], [4.3, 11.8], "表11 视觉基线")
    add_bullets(doc, ["各层级光照方向一致；屋顶、城墙、道路比例保持统一。", "气候区可有受控色差，但各层不得像完全不同的游戏。", "可以借鉴战略地图的信息组织逻辑，不直接复制现有游戏的具体资产或UI。"])

    add_heading(doc, 1, "24. 美术与程序交接")
    add_bullets(doc, ["源文件PSB／PSD、预览PNG、透明对象PNG", "地图坐标与对象锚点表、资产清单、状态列表、LOD规则", "版本记录、已知问题、交付检查表"])
    add_rule(doc, "冲突处理：画师不得自行移动对象，程序不得猜测坐标。由策划、地图负责人和美术负责人共同确认，并同步更新Excel与Word版本。")
    doc.add_paragraph("程序不得从图片颜色反向猜测势力、所有权、道路等级、通行状态、建筑状态、情报状态、商品数量或人口数量；这些事实必须来自结构化数据。")

    doc.add_page_break()
    add_heading(doc, 1, "25. 中山连续地图样板要求")
    add_numbered(doc, ["东汉天下图中的中山位置", "冀州—中山区域图", "中山县域乡野图", "中山城镇空间图", "中山商号主堂院落图"])
    add_heading(doc, 2, "25.1 连续锚点")
    add_bullets(doc, ["同一条主要河流与同一组山岭", "同一条主要官道与中山城位置", "城门方向与城外道路方向", "商号所在街区与入口方向", "仓库和车马场的相对位置"])
    add_heading(doc, 2, "25.2 禁止项")
    add_bullets(doc, ["不得独立生成五张无关图片，或为构图随意移动河流。", "不得在下一层重做道路拓扑、改变城门方向或移动商号街区。", "不得以功能卡片代替城镇空间，不得在建筑近景改变入口朝向。"])
    add_rule(doc, "考证边界：当前“中山国节点”仍是区域级代理点。具体年代、城址、水系与行政口径未确认前，不得把示意坐标写成历史定论。", True)
    add_figure_placeholder(doc, 3, "天下图至中山院落连续放大示意")

    add_heading(doc, 1, "26. 质量检查与验收")
    checks = [
        ("地理连续性", "河流流向、山脉走势、城市位置、道路连接、桥港临水、城门与入口方向"),
        ("视觉统一性", "光照、屋顶与城墙比例、道路等级、色彩、视角、阴影、材质语言"),
        ("技术拆分", "动态对象、旗帜、所有权、人物、货物、火灾、损毁、阴影可独立；文字未画死"),
        ("数据关联", "对象ID、资产ID、锚点、LOD、状态编码、版本与坐标均与Excel一致"),
    ]
    add_table(doc, ["验收域", "检查内容"], checks, [4.0, 12.1], "表12 质量验收总表")

    add_heading(doc, 1, "27. 待确认事项")
    pending = ["游戏正式起始年份", "东汉行政区采用的资料版本", "地图覆盖边界", "历史准确度与游戏化比例", "中山样板具体年代", "中山城址采用口径", "中山周边水系采用口径", "天下母版最终分辨率", "游戏引擎", "2D或2.5D表现形式", "单位移动粒度", "战场展开方式", "最终LOD切换阈值", "是否采用真实经纬度投影", "是否需要昼夜变化", "是否需要四季状态"]
    add_table(doc, ["编号", "待确认事项", "决策人", "截止日期", "结论／依据"], [(i + 1, item, "", "", "") for i, item in enumerate(pending)], [1.2, 6.8, 2.7, 2.4, 3.0], "表13 待确认事项")

    doc.add_page_break()
    add_heading(doc, 1, "附录A：标准图层树")
    add_table(doc, ["英文组名", "中文组名", "用途／导出规则"], [(a, b, c) for a, b, c in layer_tree], [4.3, 3.0, 8.8], "表A-1 标准图层树（完整）")

    add_heading(doc, 1, "附录B：标准命名示例")
    naming = [
        ("地形", "terrain_north_plain_base_v001.psb"), ("河流", "river_yellow_main_v001.png"),
        ("道路", "road_zhongshan_official_normal_l03_v001.png"), ("城市", "city_zhongshan_normal_l01_v001.png"),
        ("建筑", "building_shop_zhongshan_normal_l04_v001.png"), ("人物", "person_merchant_generic_l06_v001.png"),
        ("军队", "army_han_field_l01_v001.png"), ("商队", "caravan_horse_medium_l03_v001.png"),
        ("船只", "ship_river_cargo_l02_v001.png"), ("状态覆盖", "overlay_fire_medium_v001.png"),
        ("UI参考", "ui_reference_selection_l04_v001.png"),
    ]
    add_table(doc, ["类别", "示例"], naming, [4.0, 12.1], "表B-1 命名示例")

    add_heading(doc, 1, "附录C：资产交付检查表")
    delivery_checks = ["源文件与预览图齐全", "透明对象与阴影独立", "命名、资产ID、对象ID正确", "画布、裁切和锚点一致", "LOD与状态登记完整", "透明边缘无白边黑边", "文字、旗帜、所有权未画死", "来源、许可、版本和修改状态已登记", "已知问题已记录", "Excel与Word版本一致"]
    add_table(doc, ["检查项", "是否通过", "问题说明", "责任人", "日期"], [(x, "□", "", "", "") for x in delivery_checks], [6.2, 2.0, 4.0, 2.0, 1.9], "表C-1 资产交付检查表")

    add_heading(doc, 1, "附录D：连续性验收表")
    headers = ["验收ID", "对象ID", "检查类型", "上级地图", "下级地图", "上级X/Y", "下级换算X/Y", "偏差／允许偏差", "方向一致", "连接一致", "外形／状态／情报／损毁继承", "结果", "问题／责任／日期／复核／备注"]
    add_table(doc, headers, [("",) * len(headers) for _ in range(4)], [1.2, 1.2, 1.5, 1.35, 1.35, 1.35, 1.5, 1.55, 1.0, 1.0, 2.2, 1.0, 1.85], "表D-1 连续性验收表（字段与Excel 09_连续性验收一致；合并显示字段录入时需拆回原列）")
    doc.add_paragraph("Excel原字段完整顺序：验收ID、对象ID、检查类型、上级地图、下级地图、上级地图坐标X、上级地图坐标Y、下级换算坐标X、下级换算坐标Y、坐标偏差、允许偏差、方向一致、连接关系一致、外形继承、状态继承、情报继承、损毁继承、检查结果、问题说明、责任人、检查日期、复核人、复核日期、备注。")

    add_heading(doc, 2, "附录说明：来源与许可")
    doc.add_paragraph("正式外部素材或数据必须记录名称与版本、作者或发布机构、原始页面、下载日期、许可证及版本、修改状态、项目位置、署名要求和是否可随公开仓库再分发。禁止使用商业游戏提取资产；AI辅助概念图不得作为历史地理证据。")
    return doc


def validate(path: Path) -> None:
    errors = []
    if path.name != "03_统一地图美术规范.docx":
        errors.append("输出文件名错误")
    try:
        doc = Document(path)
    except Exception as exc:
        raise RuntimeError(f"DOCX重新打开失败: {exc}") from exc
    texts = [p.text for p in doc.paragraphs]
    texts.extend(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    joined = "\n".join(texts)
    required = [
        "统一地图美术规范", "Unified World Map Art Specification", VERSION, DOC_ID, EXCEL_NAME,
        *[f"{i}. " for i in range(1, 28)], "附录A", "附录B", "附录C", "附录D", "待项目确认图",
    ]
    for item in required:
        if item not in joined:
            errors.append(f"缺少必需内容: {item}")
    required_styles = {"标题", "副标题", "标题1", "标题2", "标题3", "正文", "引用", "图题", "表题", "警告", "规则"}
    actual_styles = {s.name for s in doc.styles}
    if missing := sorted(required_styles - actual_styles):
        errors.append(f"缺少样式: {', '.join(missing)}")
    if len(doc.sections) < 2:
        errors.append("封面未独立分节")
    if len(doc.tables) < 18:
        errors.append(f"表格数量不足: {len(doc.tables)}")
    with ZipFile(path) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
        settings_xml = zf.read("word/settings.xml").decode("utf-8")
        if "TOC" not in document_xml:
            errors.append("自动目录域不存在")
        if "updateFields" not in settings_xml:
            errors.append("未设置打开时更新域")
        if "headerReference" not in document_xml or "footerReference" not in document_xml:
            errors.append("正文页眉页脚不存在")
    if errors:
        raise RuntimeError("质量检查失败:\n- " + "\n- ".join(errors))


def main() -> int:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    validate(OUTPUT)
    print(f"Generated and validated: {OUTPUT}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(str(exc), file=sys.stderr)
        raise SystemExit(1)
